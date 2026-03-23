"""
Geração Automatizada de Documento Técnico (DEMO)

Este módulo demonstra a lógica de geração de documentos técnicos
a partir de dados estruturados.

A versão completa inclui:
- Leitura avançada de planilhas
- Regras específicas de domínio
- Geração real de arquivos Word e PDF
"""

import pandas as pd
from docx import Document
from docx.shared import Pt
import os


# ================== CONFIGURAÇÕES DEMO ==================

BASE_DIR = os.path.dirname(os.path.abspath(__file__))
ARQUIVO_EXCEL = os.path.join(BASE_DIR, "dados_exemplo.xlsx")
MODELO_DOC = os.path.join(BASE_DIR, "modelo_exemplo.docx")
PASTA_SAIDA = os.path.join(BASE_DIR, "saidas_demo")


# ================== FUNÇÕES ==================

def carregar_dados():
    """
    Carrega dados fictícios para demonstração.
    """
    dados = {
        "ID": ["01/2025", "02/2025", "03/2025"],
        "DESCRICAO": [
            "Item de demonstração A",
            "Item de demonstração B",
            "Item de demonstração C"
        ],
        "DATA": ["01/01/2025", "05/01/2025", "10/01/2025"],
        "VALOR": [1500.00, 2300.50, 999.99]
    }
    return pd.DataFrame(dados)


def gerar_documento(dados):
    """
    Gera um documento técnico simples baseado em dados estruturados.
    """
    doc = Document()

    titulo = doc.add_heading("Documento Técnico (DEMO)", level=1)
    for run in titulo.runs:
        run.font.name = "Calibri"
        run.font.size = Pt(16)

    doc.add_paragraph(
        "Este documento foi gerado automaticamente a partir de dados estruturados."
    )

    tabela = doc.add_table(rows=1, cols=4)
    hdr_cells = tabela.rows[0].cells
    hdr_cells[0].text = "ID"
    hdr_cells[1].text = "Descrição"
    hdr_cells[2].text = "Data"
    hdr_cells[3].text = "Valor"

    for _, row in dados.iterrows():
        cells = tabela.add_row().cells
        cells[0].text = row["ID"]
        cells[1].text = row["DESCRICAO"]
        cells[2].text = row["DATA"]
        cells[3].text = f"R$ {row['VALOR']:,.2f}"

    return doc


def salvar_documento(doc, nome="documento_demo.docx"):
    os.makedirs(PASTA_SAIDA, exist_ok=True)
    caminho = os.path.join(PASTA_SAIDA, nome)
    doc.save(caminho)
    print(f"[DEMO] Documento salvo em: {caminho}")


# ================== EXECUÇÃO ==================

if __name__ == "__main__":
    dados = carregar_dados()
    documento = gerar_documento(dados)
    salvar_documento(documento)
